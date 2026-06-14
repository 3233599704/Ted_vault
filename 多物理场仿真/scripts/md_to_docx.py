"""将 wiki markdown 转为 docx，去掉 Obsidian 特有语法"""
import os, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

VAULT = Path(r"D:\Staid\app\Obsidian\Ted_vault")
OUT = VAULT / "多物理场仿真" / "outputs" / "周报附件"

FILES = [
    "多物理场仿真/wiki/软件操作/WavEDA/WavEDA 索引.md",
    "多物理场仿真/wiki/软件操作/COMSOL/2026-06-09 COMSOL 6.4 安装.md",
    "多物理场仿真/wiki/软件操作/COMSOL/2026-06-10 微带贴片天线仿真.md",
    "多物理场仿真/wiki/模型训练/电磁场代理模型.md",
    "多物理场仿真/wiki/模型训练/微带天线 Port_S2 代理模型.md",
    "多物理场仿真/wiki/模型训练/dB域与线性域对比研究.md",
    "多物理场仿真/wiki/模型训练/损失函数设计.md",
    "多物理场仿真/wiki/模型训练/RAG 检索策略对比.md",
    "多物理场仿真/wiki/自动化/飞书 Bot 构建经验.md",
    "多物理场仿真/wiki/模型训练/代理模型与降阶模型.md",
    "多物理场仿真/wiki/模型训练/PINN.md",
    "多物理场仿真/wiki/索引.md",
]

def clean_md(text: str) -> str:
    """把 Obsidian markdown 转成干净的纯文本"""
    # 去掉 YAML frontmatter
    text = re.sub(r'^---\n.*?\n---\n', '', text, flags=re.DOTALL)
    # 去掉 wiki 链接 → 只保留标题
    text = re.sub(r'\[\[([^\]|#]+?)(?:#[^\]|]*?)?(?:\|[^\]]+?)?\]\]', r'\1', text)
    # 去掉图片嵌入
    text = re.sub(r'!\[\[.*?\]\]', '[图片]', text)
    # 去掉 html 注释
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)
    # 去掉 Dataview 块
    text = re.sub(r'```dataview.*?```', '', text, flags=re.DOTALL)
    # 去掉 Mermaid 块
    text = re.sub(r'```mermaid.*?```', '[流程图]', text, flags=re.DOTALL)
    # 去掉残留的 > 引用符号
    text = re.sub(r'^> ?', '', text, flags=re.MULTILINE)
    # 代码块保留内容
    text = re.sub(r'```\w*\n(.*?)```', r'\n[代码]\n\1\n', text, flags=re.DOTALL)
    # 行内代码
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # 加粗 → 普通
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # 斜体 → 普通
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    # 多余空行
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def markdown_to_lines(md_text: str) -> list:
    """把 markdown 转成 docx 段落列表"""
    lines = []
    for line in md_text.split('\n'):
        line = line.strip()
        if not line:
            lines.append(('empty', ''))
        elif line.startswith('# ') or line.startswith('## ') or line.startswith('### '):
            level = line.count('#')
            lines.append(('heading', line.lstrip('# ').strip()))
        elif line.startswith('|'):
            lines.append(('table', line))
        elif line.startswith('- ') or line.startswith('* '):
            lines.append(('bullet', line[2:]))
        else:
            lines.append(('text', line))
    return lines

os.makedirs(OUT, exist_ok=True)

for md_path in FILES:
    full = VAULT / md_path
    if not full.exists():
        print(f"SKIP: {md_path}")
        continue

    text = full.read_text(encoding='utf-8')
    text = clean_md(text)
    para_lines = markdown_to_lines(text)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(11)

    for kind, content in para_lines:
        if kind == 'empty':
            doc.add_paragraph('')
        elif kind == 'heading':
            h = doc.add_heading(content, level=1)
        elif kind == 'text':
            p = doc.add_paragraph(content)
        elif kind == 'bullet':
            p = doc.add_paragraph(content, style='List Bullet')
        elif kind == 'table':
            # 简单表格：跳过表格，转成文本
            cells = [c.strip() for c in content.split('|') if c.strip()]
            if cells and not all(c.startswith(':') or c.startswith('-') for c in cells):
                p = doc.add_paragraph(' | '.join(cells))

    out_name = Path(md_path).stem + '.docx'
    out_path = OUT / out_name
    doc.save(str(out_path))
    print(f"OK: {out_name}")

print(f"\nDone! Files in: {OUT}")
